#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JianDu Xiao Zhu Shou - Deepseek API Edition
=============================================
Flow: AI expand keywords -> Excel rough select -> AI curate books
   -> AI write article -> AI polish -> Word + HTML output

Usage:
  python jiandu_deepseek.py --theme "端午节" --max 8

Env vars:
  DEEPSEEK_API_KEY   - Deepseek API Key (required)
  DEEPSEEK_BASE_URL  - API endpoint (default https://api.deepseek.com/v1)
  DEEPSEEK_MODEL     - Model name (default deepseek-v4-pro)
"""

import argparse
import base64
import json
import os
import re
import sys
import time

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from playwright.sync_api import sync_playwright

# ── Config ──────────────────────────────────────────────
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-v4-pro")

HEADER_TEXT = "国潮汉风、快哉荐读"
FOOTER_TEXT = "书香徐来、开卷快哉"
LIBRARY_NAME = "徐州图书馆"


# ── System Prompts ──────────────────────────────────────

KEYWORD_EXPAND_SYSTEM = """\
You are a librarian helping to find books in a Chinese public library catalog \
(30,000+ titles spanning all subjects). Given a theme, generate a list of \
Chinese keywords for searching the catalog.

Important:
- The theme may refer to an upcoming holiday (e.g. Dragon Boat Festival, \
Mid-Autumn Festival), a historical commemoration, or a current social topic.
- Understand the SEMANTIC MEANING of the theme, not just the literal characters.
- Generate keywords covering: the theme itself, related cultural concepts, \
historical figures, customs, objects, philosophical ideas, artistic expressions, \
related places, and relevant time periods.
- Each keyword will be used for a substring match against book titles/authors/ \
subjects, so use specific terms that would appear in book metadata.
- Output ONLY a JSON array of strings, nothing else. Example: ["keyword1","keyword2",...]
- Generate 8-15 keywords."""

BOOK_CURATION_SYSTEM = """\
You are a librarian at Xuzhou Library. Your task is to select books from a \
candidate list that BEST match the given theme.

Rules:
- The theme should be interpreted broadly: consider upcoming holidays, \
cultural commemorations, historical events, and social context.
- Select books that are TOPICALLY relevant to the theme, NOT just keyword matches.
- Reject books that only superficially match (e.g. a book with "festival" in \
the title but about political congresses is NOT relevant to Dragon Boat Festival).
- Prefer books with cultural, literary, historical, or artistic connections.
- Prefer books that are accessible to general public readers.
- Output ONLY a JSON array of indices (0-based) of the selected books, \
sorted by relevance (most relevant first).
- Select no more than the requested maximum number of books.
- If fewer books are truly relevant, output only those. Do NOT pad with irrelevant books.
- Format: [0, 3, 7, ...]"""

RECOMMEND_SYSTEM = """\
You are a senior editor for the "KuaiZai JianDu" book recommendation column at \
Xuzhou Library, writing for general public readers.

Your article MUST follow this structure:

## 卷首语
A 2-3 paragraph opening essay that:
- Connects the theme to the joy of reading and cultural life
- Is warm, engaging, and sets the mood for the book recommendations
- Uses concise, approachable language suitable for all readers

## 推荐书籍
For each book, use this exact format:

### 1. 《书名》
- **作者**：author
- **索书号**：call_no
- **馆藏位置**：徐州图书馆参考咨询室
- **荐读理由**：150-300 Chinese characters. Write with warmth and humanity. \
Connect the book to readers' daily lives and spiritual world. \
Avoid academic or lecturing tone. Each recommendation should feel personal and inviting.

(Repeat for each book)

## 结语
A 1-2 paragraph closing that:
- Ties back to the opening theme
- Includes the slogan: "书香徐来、开卷快哉"
- Mentions Xuzhou Library
- Encourages readers to visit and borrow the featured books

Style: concise, approachable, inclusive. Elegant but not pretentious. \
Use Chinese-style book title marks 《》. \
Output ONLY the final article, no extra commentary."""

POLISH_SYSTEM = """\
You are a copy editor for Xuzhou Library. Polish this recommendation draft:

1. Fix typos, punctuation, grammar issues.
2. Improve sentence fluency and readability.
3. Preserve the structure: 卷首语 -> 推荐书籍 -> 结语.
4. Preserve ALL metadata fields (作者, 索书号, 馆藏位置).
5. Keep the warm, approachable style.
6. Output ONLY the polished full text. Do NOT explain changes."""

HTML_SYSTEM = """\
You are a frontend designer creating promotional web pages for a public library.

Generate a COMPLETE, self-contained HTML file based on the book recommendation article. Requirements:

1. **Single file**: All CSS inline, no external resources.
2. **Responsive**: Works on mobile (WeChat in-app browser) and desktop.
3. **Design**: Choose colors/fonts based on the theme atmosphere:
   - Cultural/historical: warm browns, ochre, gold, serif fonts
   - Festivals/celebrations: vibrant but elegant, festive colors appropriate to the holiday
   - Literary/arts: refined greens, slate blues, elegant serif
   - Do NOT use a cookie-cutter template.
4. **Structure**: Preserve the article structure exactly:
   - Library logo at top: use <img id="library-logo" src="LOGO_DATA_URI" alt="馆标"> exactly as-is
   - Column title (国潮汉风、快哉荐读 -- theme)
   - 卷首语 (styled as elegant intro text)
   - 推荐书籍 section (each book: title in bold, metadata line, recommendation text)
   - 结语 (closing section with slogan)
   - Library name at bottom
5. **NO emoji** anywhere. Keep it professional and elegant.
6. **Output ONLY raw HTML source code**. No ```html wrapper, no explanation.
7. **Text content must be identical to the article provided.** Do not add, remove, or \
rewrite the recommendation text."""

POSTER_SYSTEM = """\
You are a graphic designer creating a promotional poster for a public library event.

Generate a COMPLETE, self-contained HTML file optimized as a 1080x1350 pixel poster \
(for social media / printing). Requirements:

1. **Dimensions**: Fixed 1080px wide x 1350px tall. Use CSS to enforce this exact size.
2. **Design style**: Choose colors, fonts, and decorative elements based on the theme \
atmosphere. Make it visually striking and professional. Do NOT use a template.
3. **Content**: Based on the article, include:
   - Library logo image at top: use <img id="library-logo" src="LOGO_DATA_URI" alt="馆标"> exactly as-is
   - Column title: "国潮汉风、快哉荐读"
   - Theme line: the article's topic
   - A brief, elegant excerpt from the opening essay (1-2 key sentences)
   - Book list: each book with title + author (no long recommendation text - too much for a poster)
   - Slogan: "书香徐来、开卷快哉" at bottom
   - Library name: "徐州图书馆"
4. **Typography**: Use system Chinese fonts (SimSun, Microsoft YaHei, or similar). \
Title should be large and bold (60-80px). Book list should be readable at 28-36px.
5. **NO emoji** anywhere.
6. **Output ONLY raw HTML source code**. No ```html wrapper, no explanation.
7. **Design for impact**: This is a poster people will see on their phone screen or \
printed. Use visual hierarchy, negative space, and decorative elements (lines, \
geometric shapes, subtle patterns) to make it beautiful."""


# ── Excel ───────────────────────────────────────────────

def load_excel(path):
    df = pd.read_excel(path, dtype=str, header=None)
    df.fillna("", inplace=True)
    if len(df) > 4:
        df.columns = df.iloc[4].tolist()
        df = df.iloc[5:].reset_index(drop=True)
    else:
        df.columns = [f"col_{i}" for i in range(df.shape[1])]
    rename_map = {}
    for c in df.columns:
        cs = str(c)
        if re.search(r"索书|索書|call", cs, re.I):
            rename_map[c] = "call_no"
        elif re.search(r"题名|书名|title", cs, re.I):
            rename_map[c] = "title"
        elif re.search(r"责任|作者|author", cs, re.I):
            rename_map[c] = "author"
    df.rename(columns=rename_map, inplace=True)
    if "title" in df.columns:
        df = df.drop_duplicates(subset="title").reset_index(drop=True)
    return df


def rough_select(df, keywords, top_n=60):
    """Score all books by keyword match count, return top N candidates."""
    scored = []
    for idx, row in df.iterrows():
        text = " ".join(str(v) for v in row.values)
        s = sum(1 for k in keywords if k.lower() in text.lower())
        if s > 0:
            scored.append((s, idx, row.to_dict()))
    scored.sort(key=lambda x: x[0], reverse=True)
    candidates = []
    seen = set()
    for s, idx, row in scored:
        title = str(row.get("title", "")).strip()
        if not title or title in seen:
            continue
        seen.add(title)
        candidates.append({
            "title": title,
            "author": str(row.get("author", "")).strip(),
            "call_no": str(row.get("call_no", "")).strip(),
            "score": s,
        })
        if len(candidates) >= top_n:
            break
    return candidates


# ── Deepseek API ────────────────────────────────────────

def _build_client():
    if not DEEPSEEK_API_KEY:
        print("Please set DEEPSEEK_API_KEY environment variable")
        sys.exit(1)
    return OpenAI(base_url=DEEPSEEK_BASE_URL, api_key=DEEPSEEK_API_KEY)


def _call_llm(client, system, user, max_tokens=4096, temperature=0.7):
    resp = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        max_tokens=max_tokens,
        temperature=temperature,
    )
    return resp.choices[0].message.content or ""


def expand_keywords(client, theme):
    """Use AI to expand the theme into semantically relevant search keywords."""
    user = (
        f"Theme: {theme}\n\n"
        "Consider: is this an upcoming holiday? A cultural commemoration? "
        "A social hot topic? Generate search keywords that would find "
        "GENUINELY RELEVANT books in a library catalog.\n"
        "Output the JSON array now."
    )
    resp = _call_llm(client, KEYWORD_EXPAND_SYSTEM, user, max_tokens=512, temperature=0.5)
    # Parse JSON array from response
    resp = resp.strip()
    # Remove markdown code fences if present
    if resp.startswith("```"):
        resp = re.sub(r"^```\w*\n?", "", resp)
        resp = re.sub(r"\n?```$", "", resp)

    def _try_parse(text):
        try:
            result = json.loads(text)
            if isinstance(result, list):
                return [str(k) for k in result]
        except (json.JSONDecodeError, TypeError):
            pass
        return None

    # 1. Direct parse
    result = _try_parse(resp)
    if result:
        return result

    # 2. Find JSON array with greedy match
    m = re.search(r"\[.*\]", resp, re.DOTALL)
    if m:
        result = _try_parse(m.group())
        if result:
            return result

    # 3. Fallback: extract quoted strings
    kws = re.findall(r'"([^"]+)"', resp)
    if kws:
        return kws

    # 4. Last resort: split by delimiters and clean
    return [k.strip().strip('"[]\' ') for k in re.split(r"[,\n、，]", resp) if k.strip()][:15]


def ai_curate_books(client, candidates, theme, max_books):
    """Use AI to select the most relevant books from candidate list."""
    # Build a compact candidate list for the AI
    cand_lines = []
    for i, c in enumerate(candidates):
        cand_lines.append(
            f"[{i}] 《{c['title']}》| {c['author']} | {c['call_no']}"
        )
    user = (
        f"Theme: {theme}\n\n"
        f"Select the {max_books} MOST relevant books from these candidates:\n\n"
        + "\n".join(cand_lines)
        + "\n\nOutput ONLY a JSON array of selected indices, e.g. [0, 5, 12, ...]"
    )
    resp = _call_llm(client, BOOK_CURATION_SYSTEM, user, max_tokens=512, temperature=0.3)
    # Parse JSON array
    try:
        indices = json.loads(resp)
        if isinstance(indices, list):
            return [candidates[i] for i in indices if i < len(candidates)]
    except json.JSONDecodeError:
        m = re.search(r"\[.*?\]", resp, re.DOTALL)
        if m:
            try:
                indices = json.loads(m.group())
                if isinstance(indices, list):
                    return [candidates[i] for i in indices if i < len(candidates)]
            except json.JSONDecodeError:
                pass
    # Fallback: take top N by score
    return candidates[:max_books]


def generate_recommendations(client, books, theme):
    """Generate the full recommendation article in the structured format."""
    book_list = []
    for i, b in enumerate(books, 1):
        book_list.append(
            f"{i}. 《{b['title']}》\n"
            f"   作者：{b['author']}\n"
            f"   索书号：{b['call_no']}"
        )

    user = (
        f"Topic: {theme}\n\n"
        f"Selected {len(books)} books to feature:\n\n"
        + "\n\n".join(book_list)
        + "\n\nWrite the full recommendation article following the structure: "
        "卷首语 -> 推荐书籍 -> 结语."
    )
    return _call_llm(client, RECOMMEND_SYSTEM, user, max_tokens=16384, temperature=0.7)


def polish_review(client, draft, theme):
    user = f"Polish this draft for topic \"{theme}\":\n\n{draft}"
    return _call_llm(client, POLISH_SYSTEM, user, max_tokens=16384, temperature=0.3)


def generate_html_design(client, draft, theme, logo_data_uri):
    user = (
        f"Topic: {theme}\n"
        f"Column Header: {HEADER_TEXT}\n"
        f"Slogan: {FOOTER_TEXT}\n"
        f"Library: {LIBRARY_NAME}\n"
        f"Logo: use <img id=\"library-logo\" src=\"LOGO_DATA_URI\"> at top of page "
        f"(the renderer will inject the actual image).\n\n"
        f"=== ARTICLE (use this exact text, do not modify) ===\n\n{draft}\n\n"
        "Generate the complete HTML page now. Preserve ALL text exactly as written."
    )
    html = _call_llm(client, HTML_SYSTEM, user, max_tokens=24576, temperature=0.6)
    if logo_data_uri and "LOGO_DATA_URI" in html:
        html = html.replace("LOGO_DATA_URI", logo_data_uri)
    return html


def generate_poster_html(client, draft, theme, logo_path, logo_data_uri):
    """AI designs a poster layout as HTML (1080x1350)."""
    article_snapshot = draft[:1200]
    user = (
        f"Topic: {theme}\n"
        f"Column Header: {HEADER_TEXT}\n"
        f"Slogan: {FOOTER_TEXT}\n"
        f"Library: {LIBRARY_NAME}\n\n"
        f"=== Article (use as content reference) ===\n\n{article_snapshot}\n\n"
        "Design the poster HTML now. Make it beautiful and impactful.\n"
        "IMPORTANT: Include the library logo as <img id=\"library-logo\" src=\"LOGO_DATA_URI\"> "
        "(the renderer will inject the actual image data before rendering)."
    )
    html = _call_llm(client, POSTER_SYSTEM, user, max_tokens=16384, temperature=0.5)
    if logo_data_uri and "LOGO_DATA_URI" in html:
        html = html.replace("LOGO_DATA_URI", logo_data_uri)
    return html


def render_poster(html_text, out_path, width=1080, height=1350):
    """Render poster HTML to PNG using headless Chromium."""
    html_clean = html_text.strip()
    for prefix in ["```html", "```"]:
        if html_clean.startswith(prefix):
            html_clean = html_clean[len(prefix):]
    for suffix in ["```"]:
        if html_clean.endswith(suffix):
            html_clean = html_clean[:-len(suffix)]
    html_clean = html_clean.strip()
    if not html_clean.startswith("<!") and not html_clean.lower().startswith("<html"):
        html_clean = f"<!DOCTYPE html>\n{html_clean}"

    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page(viewport={"width": width, "height": height})
        page.set_content(html_clean)
        # Wait for fonts/images to load
        page.wait_for_timeout(1500)
        page.screenshot(path=out_path, full_page=False, type="png")
        browser.close()
    print(f"  Poster saved: {out_path}")


# ── Word Export ─────────────────────────────────────────

def article_to_docx(article_text, theme, out_path, logo_path=None):
    """Convert the structured article to a formatted Word document.
    Handles the 卷首语 / 推荐书籍 / 结语 structure."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.6

    # ── Logo ──
    if logo_path and os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Cm(3), height=Cm(3))
        logo_para.paragraph_format.space_after = Pt(8)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{HEADER_TEXT} —— {theme}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 50, 30)
    title.paragraph_format.space_after = Pt(20)

    # ── Parse the article into sections ──
    lines = article_text.split("\n")
    current_section = None  # "intro", "books", "closing"
    current_book = None

    def add_body_para(text, indent=False):
        if not text.strip():
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.7)
        run = p.add_run(text.strip())
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(50, 50, 50)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text.strip())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(80, 50, 30)

    def add_book_meta(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(110, 110, 110)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Detect section headers
        if stripped.startswith("## ") or stripped.startswith("# "):
            section_name = stripped.replace("#", "").strip()
            if "卷首" in section_name or "导语" in section_name or "开篇" in section_name:
                current_section = "intro"
                add_section_header("卷首语")
            elif "推荐" in section_name or "书" in section_name or "荐" in section_name:
                current_section = "books"
                add_section_header("推荐书籍")
            elif "结语" in section_name or "结尾" in section_name or "尾声" in section_name:
                current_section = "closing"
                add_section_header("结语")
            continue

        # Book title: ### N. 《书名》 or similar
        book_match = re.match(r"^(?:###\s*)?(\d+)[\.\、\)）]?\s*《(.+?)》", stripped)
        if book_match:
            num = book_match.group(1)
            btitle = book_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{num}. 《{btitle}》")
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 30, 30)
            current_section = "books"
            continue

        # Book metadata lines
        for field in ["作者", "索书号", "馆藏位置"]:
            if stripped.startswith(f"- **{field}**") or stripped.startswith(f"- {field}") or stripped.startswith(f"{field}"):
                val = re.sub(r"^-?\s*\*{0,2}" + field + r"\*{0,2}[：:]\s*", "", stripped)
                add_book_meta(field, val)
                break
        else:
            # Check if it's a 荐读理由 line
            if stripped.startswith("- **荐读理由**") or stripped.startswith("- 荐读理由"):
                continue  # header line, skip
            # Regular paragraph
            if current_section == "books":
                add_body_para(stripped, indent=True)
            elif current_section in ("intro", "closing"):
                add_body_para(stripped, indent=True)
            else:
                add_body_para(stripped)

    # ── Footer ──
    footer_div = doc.add_paragraph()
    footer_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_div.paragraph_format.space_before = Pt(30)
    run = footer_div.add_run("—" * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(160, 140, 110)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(FOOTER_TEXT)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 50, 30)
    footer.paragraph_format.space_after = Pt(4)

    lib = doc.add_paragraph()
    lib.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lib.add_run(LIBRARY_NAME)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(130, 120, 100)

    doc.save(out_path)
    print(f"  Word saved: {out_path}")


# ── Main ────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="JianDu - Deepseek API")
    parser.add_argument("--excel", default="参考咨询书库馆藏清单.xlsx")
    parser.add_argument("--theme", required=True)
    parser.add_argument("--max", type=int, default=8)
    parser.add_argument("--out_dir", default="output")
    parser.add_argument("--logo", default="馆标黑版.png", help="Library logo PNG")
    parser.add_argument("--skip_ai", action="store_true", help="Skip LLM calls (debug mode)")
    args = parser.parse_args()

    os.makedirs(args.out_dir, exist_ok=True)

    banner = "=" * 60
    print(banner)
    print(f"  JianDu Xiao Zhu Shou")
    print(f"  Topic: {args.theme}")
    print(banner)

    if not os.path.exists(args.excel):
        print(f"Not found: {args.excel}")
        sys.exit(1)

    if args.skip_ai:
        # ── Debug mode: keyword only ──
        kws = re.split(r"[,,\s]+", args.theme.strip())
        kws = [k for k in kws if k]
        df = load_excel(args.excel)
        print(f"Loaded {len(df)} records")
        books = rough_select(df, kws, args.max)
        print(f"Keyword match: {len(books)} books")
        for i, b in enumerate(books):
            print(f"  {i+1}. 《{b['title']}》 -- {b['author']}")
        sys.exit(0)

    # ── Full AI flow ──
    client = _build_client()

    # 1. Expand keywords
    print("\n[1/6] AI expanding keywords from theme...")
    keywords = expand_keywords(client, args.theme)
    # Merge user's original theme words
    user_kws = re.split(r"[,,\s]+", args.theme.strip())
    user_kws = [k for k in user_kws if k]
    all_keywords = list(dict.fromkeys(user_kws + keywords))  # dedupe, preserve order
    print(f"  Keywords ({len(all_keywords)}): {', '.join(all_keywords[:10])}" +
          ("..." if len(all_keywords) > 10 else ""))

    # 2. Rough select from Excel
    print("\n[2/6] Searching catalog...")
    df = load_excel(args.excel)
    print(f"  Catalog: {len(df)} records")
    candidates = rough_select(df, all_keywords, top_n=60)
    print(f"  Candidates: {len(candidates)} books matched")
    for i, c in enumerate(candidates[:10]):
        print(f"    [{i}] 《{c['title']}》 {c['author']}")
    if len(candidates) > 10:
        print(f"    ... and {len(candidates)-10} more")

    if not candidates:
        print("  No books found. Try a different theme.")
        sys.exit(0)

    logo_path = args.logo if os.path.exists(args.logo) else None
    logo_data_uri = ""
    if logo_path:
        with open(logo_path, "rb") as f:
            logo_data_uri = "data:image/png;base64," + base64.b64encode(f.read()).decode("ascii")
            print(f"  Logo encoded ({len(logo_data_uri)} chars)")

    # 3. AI curate
    print(f"\n[3/6] AI selecting best {args.max} books from candidates...")
    time.sleep(0.3)
    books = ai_curate_books(client, candidates, args.theme, args.max)
    print(f"  Selected {len(books)} books:")
    for i, b in enumerate(books, 1):
        print(f"    {i}. 《{b['title']}》 — {b['author']} ({b['call_no']})")

    if not books:
        print("  AI found no relevant books. Try a different theme.")
        sys.exit(0)

    # 4. Write article
    print("\n[4/6] AI writing recommendation article...")
    time.sleep(0.3)
    draft = generate_recommendations(client, books, args.theme)
    print("-" * 40)
    preview = draft[:600] + "..." if len(draft) > 600 else draft
    print(preview)
    print("-" * 40)

    # 5. Polish
    print("\n[5/6] AI polishing...")
    time.sleep(0.3)
    final = polish_review(client, draft, args.theme)
    print("  Polishing done")

    # 6. HTML design
    print("\n[6/6] AI generating HTML page...")
    time.sleep(0.3)
    html = generate_html_design(client, final, args.theme, logo_data_uri)
    print("  HTML generated")

    # 7. Poster
    print("\n[7/7] AI designing poster...")
    time.sleep(0.3)
    poster_html = generate_poster_html(client, final, args.theme, logo_path, logo_data_uri)
    print("  Poster HTML designed")

    # ── Save outputs ──
    safe_theme = re.sub(r"[^\w一-鿿]+", "_", args.theme.strip()).strip("_")

    # Markdown
    md_path = os.path.join(args.out_dir, f"jiandu_{safe_theme}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(final)
    print(f"\n  Markdown saved: {md_path}")

    # Word
    docx_path = os.path.join(args.out_dir, f"jiandu_{safe_theme}.docx")
    article_to_docx(final, args.theme, docx_path, logo_path)

    # HTML
    if html:
        html_path = os.path.join(args.out_dir, f"jiandu_{safe_theme}.html")
        html_clean = html.strip()
        for prefix in ["```html", "```"]:
            if html_clean.startswith(prefix):
                html_clean = html_clean[len(prefix):]
        for suffix in ["```"]:
            if html_clean.endswith(suffix):
                html_clean = html_clean[:-len(suffix)]
        html_clean = html_clean.strip()
        if not html_clean.startswith("<!") and not html_clean.lower().startswith("<html"):
            html_clean = f"<!DOCTYPE html>\n{html_clean}"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_clean)
        print(f"  HTML saved: {html_path}")

    # Poster (HTML source + rendered PNG)
    if poster_html:
        poster_html_path = os.path.join(args.out_dir, f"jiandu_{safe_theme}_poster.html")
        poster_html_clean = poster_html.strip()
        for prefix in ["```html", "```"]:
            if poster_html_clean.startswith(prefix):
                poster_html_clean = poster_html_clean[len(prefix):]
        for suffix in ["```"]:
            if poster_html_clean.endswith(suffix):
                poster_html_clean = poster_html_clean[:-len(suffix)]
        poster_html_clean = poster_html_clean.strip()
        if not poster_html_clean.startswith("<!"):
            poster_html_clean = f"<!DOCTYPE html>\n{poster_html_clean}"
        with open(poster_html_path, "w", encoding="utf-8") as f:
            f.write(poster_html_clean)
        print(f"  Poster HTML saved: {poster_html_path}")

        poster_path = os.path.join(args.out_dir, f"jiandu_{safe_theme}_poster.png")
        render_poster(poster_html_clean, poster_path)

    print("\n" + banner)
    print("Done! Output files:")
    for fn in sorted(os.listdir(args.out_dir)):
        if fn.startswith("jiandu_") and safe_theme in fn:
            fp = os.path.join(args.out_dir, fn)
            size_kb = os.path.getsize(fp) / 1024
            print(f"  {fn} ({size_kb:.1f} KB)")
    print(banner)


if __name__ == "__main__":
    main()
